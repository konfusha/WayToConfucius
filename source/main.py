from mainMenu import Ui_MainWindow
from addWordsWindow import Ui_AddWordsWindow
from trainerWindow import Ui_TrainerWindow
from statisticsWindow import Ui_Statistics
from DailyAimChoice import Ui_DailyAimChoiceWindow
from GeneratorWindow import Ui_GeneratorWindow
from ImportLoading import Ui_ImportLoadingWindow
from PyQt6.QtWidgets import (QApplication, QMainWindow, QHeaderView, QTableWidgetItem,
                             QAbstractItemView, QMessageBox, QFileDialog, QLineEdit)
from PyQt6 import QtGui
from PyQt6.QtGui import QPixmap, QImage
from PyQt6.QtCore import Qt
import sys
import sqlite3
import random as rd
import datetime
import matplotlib.pyplot as plt
import json
import docx
import time

class MainWindow(QMainWindow, Ui_MainWindow):
    def __init__(self, application: QApplication) -> None:
        super().__init__()
        self.setupUi(self)
        self.app = application
        self.setWindowTitle("Тренажёр ударений")
        
        self.startButton.clicked.connect(self.start_training)
        self.editTasksButton.clicked.connect(self.show_add_words_window)
        self.statisticsButton.clicked.connect(self.show_statistics_window)
        self.generatorButton.clicked.connect(self.show_generator_window)

    def start_training(self) -> None:
        '''
        Calls DailyAimChoice windos if it's first usage in a day.
        '''
        with open("AppData.json") as data_file:
            app_data = json.load(data_file)
        if app_data["todayDate"] != str(datetime.date.today()):
            self.dailyChoice = DailyAimChoiceWindow(self)
            self.dailyChoice.show()
        else:
            self.show_trainer_window()

    def show_trainer_window(self) -> None:
        con = sqlite3.connect("accent_tasks.sqlite")
        if len(con.execute("SELECT correct_word FROM tasks").fetchall()) == 0:
            QMessageBox.critical(self, "Ошибка",
                                 "В словаре нет слов. Добавьте слова для продолжения.")
            return
        self.trainer = AcccentTrainer(self)
        self.trainer.show()
        self.close()

    def show_add_words_window(self) -> None:
        self.add_words_window = AddWordsWindow(self)
        self.add_words_window.show()
        self.close()

    def show_statistics_window(self) -> None:
        self.statistics = StatisticsWindow(self)
        self.statistics.show()
        self.close()

    def show_generator_window(self) -> None:
        self.generator = GeneratorWindow(self)
        self.generator.show()
        self.generatorButton.setDisabled(True)


class AcccentTrainer(QMainWindow, Ui_TrainerWindow):
    def __init__(self, parent_window: MainWindow) -> None:
        super().__init__()
        self.parent_window = parent_window
        self.setupUi(self)
        self.setWindowTitle("Тренажёр ударений")
        self.correctionEdit.hide()
        self.correctionBar.hide()
        self.stat = sqlite3.connect("user_data.sqlite")
        self.tasks = sqlite3.connect("accent_tasks.sqlite")
        self.current_word_num = 0
        self.current_word = None
        self.trained_word_count = 0
        
        with open("AppData.json") as data_file:
            app_data = json.load(data_file)
            aim = app_data["todayAim"]
            self.trained_word_count = app_data["currentProgress"]
        self.globalProgressBar.setMaximum(aim)
        self.globalProgressBar.setValue(min(self.trained_word_count, aim))
        
        self.make_task()

        self.backButton.clicked.connect(self.back_to_main_window)
        self.firstCaseButton.clicked.connect(self.check_ans)
        self.secondCaseButton.clicked.connect(self.check_ans)
        self.correctionEdit.textChanged.connect(self.check_correction_word)
    

    def back_to_main_window(self) -> None:
        self.parent_window.show()
        self.close()

    def make_task(self) -> None:
        '''
        Generates a random task and saves correct pronunciation into self.correct_ans
        '''
        data = self.tasks.execute("SELECT correct_word, wrong_word FROM tasks").fetchall()
        new_word_num = rd.randint(1, len(data) - 1)
        while new_word_num == self.current_word_num:
            new_word_num = rd.randint(1, len(data) - 1)
        self.current_word_num = new_word_num


        self.current_word = data[self.current_word_num]
        self.correct_ans = self.current_word[0]
        self.current_word = list(self.current_word)
        rd.shuffle(self.current_word)

        self.firstCaseButton.setText(self.current_word[0])
        self.secondCaseButton.setText(self.current_word[1])

    def check_ans(self) -> None:
        '''
        Checks pushed button with correct answer and displays the result on the label.
        '''
        if self.sender().text() == self.correct_ans:
            self.resultLabel.setText("Верно!")
            self.correctionLabel.setText("")
            self.make_task()
            self.trained_word_count += 1
            self.globalProgressBar.setValue(self.globalProgressBar.value() + 1)
            if self.trained_word_count == self.globalProgressBar.maximum():
                QMessageBox.information(self, "Цель достигнута",
                                        "Вы успешно справились с дневной целью!")
        else:
            self.resultLabel.setText("Неверно!")
            self.count_mistake(self.correct_ans)
            self.make_correction()

    def make_correction(self) -> None:
        '''
        Called when error is made
        '''
        self.correctionLabel.setText(f"Введите правильный вариант {self.correct_ans} 3 раза")
        self.firstCaseButton.setDisabled(True)
        self.secondCaseButton.setDisabled(True)
        self.correctionEdit.show()
        self.correctionBar.show()
        
    def check_correction_word(self) -> None:
        if self.correctionEdit.text() == self.correct_ans:
            self.correctionBar.setValue(min(self.correctionBar.value() + int(100 / 3) + 1, 100))
            self.correctionEdit.clear()
        if self.correctionBar.value() == 100:
            self.trained_word_count += 1
            self.globalProgressBar.setValue(self.globalProgressBar.value() + 1)
            self.correctionBar.reset()
            self.correctionBar.hide()
            self.correctionEdit.hide()
            self.correctionLabel.clear()
            self.resultLabel.clear()
            self.firstCaseButton.setDisabled(False)
            self.secondCaseButton.setDisabled(False)
            self.make_task()
    
    def closeEvent(self, event) -> None:
        '''
        Saves completed words amount before exit
        '''
        count = self.stat.execute("SELECT count FROM daily_count WHERE date = ?",
                                  (str(datetime.date.today()),)).fetchall()
        if len(count) == 0:
            self.stat.execute("INSERT INTO daily_count(date, count) VALUES(?, ?)",
                            (str(datetime.date.today()), self.trained_word_count))
        else:
            self.stat.execute("""UPDATE daily_count
                              SET count = ?
                              WHERE date = ?""",
                              (self.trained_word_count, str(datetime.date.today())))
            
        with open("AppData.json") as data_file:
            app_data = json.load(data_file)
            app_data["currentProgress"] = self.trained_word_count
            
        with open("AppData.json", 'w') as data_file:
            json.dump(app_data, data_file, indent=4)
        self.stat.commit()
    
    def count_mistake(self, word) -> None:
        '''
        Increase amount of mistakes in data_base
        '''
        count = self.stat.execute("SELECT count FROM mistakes_count WHERE correct_word = ?",
                                  (word,)).fetchall()
        if len(count) == 0:
            self.stat.execute("INSERT INTO mistakes_count(correct_word, count) VALUES(?, ?)",
                            (word, 1))
        else:
            self.stat.execute("""UPDATE mistakes_count
                              SET count = ?
                              WHERE correct_word = ?""",
                              (count[0][0] + 1, word))
        self.stat.commit()
    


class AddWordsWindow(QMainWindow, Ui_AddWordsWindow):
    def __init__(self, parent_window: MainWindow) -> None:
        super().__init__()
        self.parent_window = parent_window
        self.setupUi(self)
        self.setWindowTitle("Редактирование слов")

        self.make_table()
        self.update_table()

        self.backButton.clicked.connect(self.back_to_main_window)
        self.addButton.clicked.connect(self.add_new_word)
        self.deleteButton.clicked.connect(self.delete_word)
        self.importButton.clicked.connect(self.import_tasks)
        self.exportButton.clicked.connect(self.export_tasks)

    def back_to_main_window(self) -> None:
        self.parent_window.show()
        self.close()

    def make_table(self) -> None:
        self.wordsTable.setColumnCount(2)
        self.wordsTable.setHorizontalHeaderLabels(
            ["Верный вариант", "Неверный вариант"])
        header = self.wordsTable.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.wordsTable.setEditTriggers(
            QAbstractItemView.EditTrigger.NoEditTriggers)

    def update_table(self) -> None:
        '''
        Updates table content
        '''
        con = sqlite3.connect("accent_tasks.sqlite")
        self.wordsTable.clearContents()
        self.wordsTable.setItem(1, 1, QTableWidgetItem("pashalco"))
        table = con.execute("SELECT * from tasks").fetchall()
        self.wordsTable.setRowCount(len(table))
        self.wordsTable.setVerticalHeaderLabels(
            [str(i)for i in range(len(table), 0, -1)])

        for i in range(len(table)):
            for j in range(2):
                self.wordsTable.setItem(
                    i, j,QTableWidgetItem(str(table[len(table) - i - 1][j + 1])))

    
    def add_new_word(self) -> None:
        con = sqlite3.connect("accent_tasks.sqlite")
        correct_w, wrong_w = self.correctEdit.text(), self.wrongEdit.text()
        try:
            if not correct_w or not wrong_w:
                raise ValueError
            if correct_w.lower() != wrong_w.lower():
                raise DifferentWords
            if not correct_w.isalpha() or not wrong_w.isalpha():
                raise WrongInput
            if len(letter1 := [x for x in correct_w if x.isupper()]) != 1 or\
                len(letter2 := [x for x in wrong_w if x.isupper()]) != 1:
                raise WrongAcccent
            not_vowel = "БВГДЖЗЙКЛМНПРСТФХЦЧШЩЪЬ"
            if letter1[0] in not_vowel or letter2[0] in not_vowel:
                raise WrongAcccent
            if correct_w == wrong_w:
                raise SamePronunciation
            correct_words = [x[0].lower() for x in
                             con.execute("SELECT correct_word FROM tasks").fetchall()]
            if correct_w.lower() in correct_words:
                raise NotUniqueWord
        except ValueError:
            self.emptyWarningLabel.setText("Введите произношение")
            return
        except DifferentWords:
            QMessageBox.critical(self, "Неверный ввод", "Введены разные слова")
            return
        except WrongInput:
            QMessageBox.critical(self, "Неверный ввод", "Введено неверное слово")
            return
        except SamePronunciation:
            QMessageBox.critical(
                self, "Неверный ввод", "Введите слово с различными варинтами ударения")
            return
        except WrongAcccent:
            QMessageBox.critical(self, "Неверный ввод", "Поставьте верное ударение")
            return
        except NotUniqueWord:
            QMessageBox.critical(self, "Неверный ввод", "Данное слово уже имеется")
            return
        self.emptyWarningLabel.setText("")
        con.execute("INSERT INTO tasks (correct_word,wrong_word) VALUES(?, ?)", (correct_w, wrong_w))
        con.commit()
        self.update_table()
        
    def delete_word(self) -> None:
        '''
        Deletes selected word(s) from data with tasks.
        '''
        con = sqlite3.connect("accent_tasks.sqlite")
        if not self.wordsTable.selectedItems():
            QMessageBox().warning(self, "Неизвестный элемент", "Выберите слова для удаления")
            return
        selected_num = sorted(set(self.wordsTable.verticalHeaderItem(
            i.row()).text() for i in self.wordsTable.selectedItems()))
        confirmBox = QMessageBox()
        valid = confirmBox.question(
            self, 'Подтвердите действие', "Действительно удалить элементы под номером " +
            ", ".join(selected_num),
            buttons=QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if valid == QMessageBox.StandardButton.No:
            return
        cur = con.cursor()
        cur.execute("DELETE FROM tasks WHERE correct_word IN (" + ", ".join('?' * len(selected_num)) + ")",
                         [self.wordsTable.item(self.wordsTable.rowCount() - int(x), 0).text() for x in selected_num])
        con.commit()   
        self.update_table()
    
    def import_tasks(self) -> None:
        fileName = QFileDialog.getOpenFileName(
            parent=self,
            caption= "Выберите название файла",
            filter= "Text files (*.docx)"
        )
        if not fileName[0]:
            return
        doc = docx.Document(fileName[0])
        message_critical = QMessageBox()
        if len(doc.paragraphs) < 2:
            message_critical.critical(self, "Неверный файл", "Откройте файл со словами, созданный этой программой")
            return
        data = [x for x in doc.paragraphs[1].text.split("\n") if x]
        if len(data) < 2:
            message_critical.critical(self, "Неверный файл", "Откройте файл со словами, созданный этой программой")
            return
        print(data)
        data = [x.split(' - ') for x in data[1:]]
        for words in data:
            try:
                if len(words) != 2:
                    raise WrongFileImported
                correct_w, wrong_w = words[0], words[1]
                if not correct_w or not wrong_w:
                    raise WrongFileImported
                if correct_w.lower() != wrong_w.lower():
                    raise WrongFileImported
                if not correct_w.isalpha() or not wrong_w.isalpha():
                    raise WrongFileImported
                if len(letter1 := [x for x in correct_w if x.isupper()]) != 1 or\
                    len(letter2 := [x for x in wrong_w if x.isupper()]) != 1:
                    raise WrongFileImported
                not_vowel = "БВГДЖЗЙКЛМНПРСТФХЦЧШЩЪЬ"
                if letter1[0] in not_vowel or letter2[0] in not_vowel:
                    raise WrongFileImported
                if correct_w == wrong_w:
                    raise WrongFileImported
            except WrongFileImported:
                message_critical.critical(self, "Неверный файл", "Откройте файл со словами, созданный этой программой")
                return
        result = QMessageBox().question(
        self, 'Подтвердите действие', "Добавить слова из файла к текущим?",
        buttons=QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if result == QMessageBox.StandardButton.No:
            return
        con = sqlite3.connect("accent_tasks.sqlite")
        correct_words_in_prev_db = [x[0] for x in con.execute("SELECT correct_word FROM tasks").fetchall()]
        sleep_time = 2 / len(data)
        self.progress_window = ImportLoadingWindow()
        self.progress_window.progressBar.setMaximum(len(data))
        self.progress_window.move(523, 132)
        self.progress_window.show()
        self.progress_window.progressBar.setValue(0)
        for words in data:
            time.sleep(sleep_time)
            self.parent_window.app.processEvents()
            self.progress_window.progressBar.setValue(self.progress_window.progressBar.value() + 1)
            correct_w, wrong_w = words[0], words[1]
            if correct_w in correct_words_in_prev_db:
                continue
            self.correctEdit.setText(correct_w)
            self.wrongEdit.setText(wrong_w)
            con.execute("INSERT INTO tasks (correct_word,wrong_word) VALUES(?, ?)", (correct_w, wrong_w))
        con.commit()
        self.progress_window.close()
        self.correctEdit.clear()
        self.wrongEdit.clear()
        QMessageBox().information(self, "Успешно", "Слова успешно добавлены")
        self.update_table()
    
    def export_tasks(self) -> None:
        fileName = QFileDialog.getSaveFileName(
            parent=self,
            caption= "Выберите название файла",
            filter= "Text files (*.docx)"
        )
        if not fileName[0]:
            return
        doc = docx.Document()
        con = sqlite3.connect("accent_tasks.sqlite")
        words =  con.execute("SELECT correct_word, wrong_word FROM tasks").fetchall()
        doc.add_paragraph("Верный вариант - неверный")
        doc.add_paragraph('\n'.join(f"{x[0]} - {x[1]}" for x in words))
        doc.save(fileName[0])
        QMessageBox.information(self, "Сохранения завершено", "Задания успешно сохранены")

class StatisticsWindow(QMainWindow, Ui_Statistics):
    def __init__(self, parent_window: MainWindow) -> None:
        super().__init__()
        self.parent_window = parent_window
        self.setupUi(self)
        self.setWindowTitle("Статистика")
        self.make_table()
        self.load_bar()
        
        with open("AppData.json") as data_file:
            word_amount = json.load(data_file)["currentProgress"]
        self.todayStatLabel.setText(f"Сегодня вы выполнили {word_amount} слов.")
        
        self.backButton.clicked.connect(self.back_to_main_window)
        self.clearHistoryButton.clicked.connect(self.clear_table)
        
    def back_to_main_window(self) -> None:
        self.parent_window.show()
        self.close()

    def make_table(self) -> None:
        con = sqlite3.connect("user_data.sqlite")
        table = con.execute("SELECT * FROM mistakes_count").fetchall()
        table.sort(key=lambda x: x[1], reverse=True)
        
        self.mistakesTable.setColumnCount(2)
        self.mistakesTable.setHorizontalHeaderLabels(["Слово", "Кол-во ошибок"])
        self.mistakesTable.setRowCount(len(table))
        header = self.mistakesTable.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        
        for i in range(len(table)):
            self.mistakesTable.setItem(i, 0, QTableWidgetItem(str(table[i][0])))
            self.mistakesTable.setItem(i, 1, QTableWidgetItem(str(table[i][1])))
        self.mistakesTable.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        
    def clear_table(self) -> None:
        '''
        Removes amount of mistakes from data
        '''
        con = sqlite3.connect("user_data.sqlite")
        confirmBox = QMessageBox()
        valid = confirmBox.question(
            self, 'Подтвердите действие', "Вы действительно хотите очистить историю?",
            buttons=QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if valid == QMessageBox.StandardButton.No:
            return
        con.execute("DELETE FROM mistakes_count")
        con.commit()
        self.make_table()
    
    def load_bar(self) -> None:
        '''
        Takes and sorts data from data_user, makes jpg with bar and displays it
        '''
        con = sqlite3.connect("user_data.sqlite")
        data = con.execute("SELECT * FROM daily_count").fetchall()
        data.sort(key=lambda x: int(x[0].split('-')[1]) * 32 + int(x[0].split('-')[2]))
        data = data[-7:]
        dates = ['-'.join(x[0].split('-')[1:]) for x in data]
        word_count = [x[1] for x in data]
        
        fig = plt.figure()
        ax = fig.add_subplot(111)
        ax.set_ylabel("Кол-во слов")
        ax.xaxis.label.set_color("white")
        ax.yaxis.label.set_color("white")
        ax.tick_params(colors='white')

        ax.spines["left"].set_color("white")
        ax.spines["top"].set_color("white") 
        ax.spines["bottom"].set_color("white") 
        ax.spines["right"].set_color("white") 
        plt.bar(dates, word_count, color="#00ffea")
        plt.gcf().set_size_inches(5, 2.5)
        plt.savefig("week_bar.png", dpi=1000, transparent=True)
        self.weeklyGraph.setPixmap(QPixmap(QImage("week_bar.png")).scaled(
            self.weeklyGraph.width(), self.weeklyGraph.height()))


class DailyAimChoiceWindow(QMainWindow, Ui_DailyAimChoiceWindow):
    def __init__(self, parent_window: MainWindow):
        self.parent_window = parent_window
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle("Ежедневный выбор")
        
        self.acceptButton.clicked.connect(self.change_daily_aim)
    
    def change_daily_aim(self) -> None:
        try:
            if not self.amountEntry.text():
                raise ValueError
            if not self.amountEntry.text().isnumeric():
                raise WrongInput
            if not (0 < int(self.amountEntry.text()) < 2**31 - 1):
                raise WrongInput
        except ValueError:
            QMessageBox.critical(self, "Неверный ввод", "Введите кол-во слов")
            return
        except WrongInput:
            QMessageBox.critical(self, "Неверный ввод", "Введено неверное число")
            return
        with open("AppData.json", 'w') as data_file:
            new_app_data = {
                "todayDate": str(datetime.date.today()),
                "currentProgress": 0,
                "todayAim": int(self.amountEntry.text())
            }
            json.dump(new_app_data, data_file, indent=4)
        self.close()
        self.parent_window.show_trainer_window()
    
    
class GeneratorWindow(QMainWindow, Ui_GeneratorWindow):
    def __init__(self, parent_window: QMainWindow) -> None:
        super().__init__()
        self.parent_window = parent_window
        self.setupUi(self)
        self.setWindowTitle("Генератор заданий")
        
        self.tasksAmountEntry.setText('10')
        self.acceptButton.clicked.connect(self.generate_text)
    
    def closeEvent(self, event) -> None:
        self.parent_window.generatorButton.setDisabled(False)
        
    def generate_text(self) -> None:
        """
        Create docx file with tasks
        """
        con = sqlite3.connect("accent_tasks.sqlite")
        words = con.execute("SELECT correct_word, wrong_word FROM tasks").fetchall()
        if len(words) < 5:
            QMessageBox.critical(self, "Недостаточно слов",
                                 "Недостаточно слов для генерации слов, добавьте хотя бы 5 слов")
            return
        if not self.tasksAmountEntry.text():
            QMessageBox.critical(self, "Неверное количество",
                                 "Введите кол-во заданий")
            return
        if (not self.tasksAmountEntry.text().isnumeric()) or int(self.tasksAmountEntry.text()) < 1:
            QMessageBox.critical(self, "Неверное количество",
                                 "Введите корректное кол-во заданий")
            return
        if self.withoutRepeatRadio.isChecked() and len(words) < int(self.tasksAmountEntry.text()) * 5:
                QMessageBox.critical(self, "Недостаточное кол-во слов",
                            f"Не хватает слов для генерации заданий {len(words)} из {int(self.tasksAmountEntry.text()) * 5}")
                return
        fileName = QFileDialog.getSaveFileName(
            parent=self,
            caption= "Выберите название файла",
            filter= "Text files (*.docx)"
        )
        if not fileName[0]:
            return
        doc = docx.Document()
        doc.add_heading("Задания на ударение", 0).add_run()
        answers = list()
        if self.repeatRadio.isChecked():
            for task_num in range(int(self.tasksAmountEntry.text())):
                rand_words_choice = rd.sample(words, k=5)
                local_ans = rd.sample(range(5), k=rd.randint(1, 4))
                answers.append(local_ans)
                par1 = doc.add_paragraph()
                par1.add_run(f"№{task_num + 1} ").bold = True
                par1.add_run("Укажите варианты ответов, в которых верно выделена буква," +\
                    "обозначающая ударный гласный звук. Запишите номера ответов.\n" +\
                        "\n".join(f"{i + 1}.) {rand_words_choice[i][0]}" if i in local_ans else
                                  f"{i + 1}.) {rand_words_choice[i][1]}" for i in range(5)))
        else:
            rand_words_choice = rd.sample(words, k=int(self.tasksAmountEntry.text()) * 5)
            for task_num in range(int(self.tasksAmountEntry.text())):
                local_rand_words = rand_words_choice[task_num * 5: (task_num + 1) * 5]
                local_ans = rd.sample(range(5), k=rd.randint(1, 4))
                answers.append(local_ans)
                par1 = doc.add_paragraph()
                par1.add_run(f"№{task_num + 1} ").bold = True
                par1.add_run("Укажите варианты ответов, в которых верно выделена буква," +\
                    "обозначающая ударный гласный звук. Запишите номера ответов.\n" +\
                        "\n".join(f"{i + 1}.) {local_rand_words[i][0]}" if i in local_ans else
                                  f"{i + 1}.) {local_rand_words[i][1]}" for i in range(5)))
                 
        doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
        doc.add_paragraph("ОТВЕТЫ:\n" + "\n".join([f"{i + 1}.) {"".join(str(x + 1) for x in answers[i])}"
                                    for i in range(len(answers))]))
        doc.save(fileName[0])
        QMessageBox.information(self, "Успешно выполнено", "Файл успешно создан")
        self.close()
        
    
class ImportLoadingWindow(QMainWindow, Ui_ImportLoadingWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle("Загрузка")

            
class WrongInput(Exception):
    pass


class DifferentWords(Exception):
    pass


class SamePronunciation(Exception):
    pass


class WrongAcccent(Exception):
    pass


class NotUniqueWord(Exception):
    pass


class WrongFileImported(Exception):
    pass
    
if __name__ == '__main__':
    app = QApplication(sys.argv)
    main_window = MainWindow(app)
    main_window.show()
    sys.exit(app.exec())
